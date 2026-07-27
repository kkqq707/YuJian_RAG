"""企业知识库文档加载器

支持格式：.txt / .md / .pdf / .docx / .xlsx
TXT / MD 编码自动检测（utf-8 → utf-8-sig → gb18030 → gbk）
每个 Document 附带完整 metadata，document_id 稳定可复现。

支持从多个来源加载:
- data/builtin: 内置知识库文件
- data/uploads: 客户上传文件
- data/: 向后兼容旧路径
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document

from src.config import DATA_DIR, BUILTIN_DATA_DIR, UPLOADS_DATA_DIR

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------------
_SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx", ".xlsx"}
_TXT_ENCODINGS = ["utf-8", "utf-8-sig", "gb18030", "gbk"]

# 临时文件 / 隐藏文件 过滤规则
_EXCLUDE_PATTERNS = re.compile(
    r"(^~\$|^\.|\.tmp$|\.temp$|\.bak$|~$)",
    re.IGNORECASE,
)

# DOCX 每个 "页" 大约的字符数（仅用于分组段落，不做真正切分）
_DOCX_PAGE_CHARS = 3000


# ---------------------------------------------------------------------------
# 公开函数
# ---------------------------------------------------------------------------


def find_knowledge_files(directory: Path | None = None) -> list[Path]:
    """递归扫描目录，返回支持的知识库文件列表（按相对路径稳定排序）。

    默认只扫描 data/uploads 目录（管理员上传的文件）。
    不再扫描 data/builtin 目录。
    如果提供 directory 参数则只扫描该目录。

    Parameters
    ----------
    directory : Path, optional
        指定扫描目录，默认只扫描 uploads

    Returns
    -------
    list[Path]
    """
    directories: list[Path] = []
    if directory is not None:
        directories = [directory]
    else:
        if UPLOADS_DATA_DIR.is_dir():
            directories.append(UPLOADS_DATA_DIR)
        # 向后兼容：如果 uploads 目录不存在，回退到 DATA_DIR
        if not directories:
            directories = [DATA_DIR]

    files: list[Path] = []
    for d in directories:
        if not d.is_dir():
            continue
        for path in sorted(d.rglob("*")):
            if path.is_file() and _is_supported(path):
                files.append(path)
    return files


def load_txt_file(file_path: Path) -> list[Document]:
    """加载 TXT 文件，依次尝试多种中文编码。

    Parameters
    ----------
    file_path : Path
        TXT 文件路径

    Returns
    -------
    list[Document]
    """
    content, detected_encoding = _read_text_with_encoding(file_path)
    doc_id = _make_document_id(file_path)
    return [
        Document(
            page_content=content,
            metadata=_make_metadata(file_path, doc_id, page=1, encoding=detected_encoding),
        )
    ]


def load_markdown_file(file_path: Path) -> list[Document]:
    """加载 Markdown 文件，编码检测同 TXT。

    Parameters
    ----------
    file_path : Path
        .md 文件路径

    Returns
    -------
    list[Document]
    """
    return load_txt_file(file_path)  # 处理逻辑完全相同


def load_pdf_file(file_path: Path) -> list[Document]:
    """加载 PDF 文件，每页生成一个 Document。

    使用 pypdf.PdfReader；某页无法提取文字时跳过并记录警告。

    Parameters
    ----------
    file_path : Path
        .pdf 文件路径

    Returns
    -------
    list[Document]
    """
    from pypdf import PdfReader

    doc_id = _make_document_id(file_path)
    documents: list[Document] = []
    try:
        reader = PdfReader(str(file_path))
    except Exception:
        logger.warning("无法打开 PDF 文件: %s", file_path, exc_info=True)
        return documents

    for i, page_obj in enumerate(reader.pages, start=1):
        try:
            text = page_obj.extract_text()
        except Exception:
            logger.warning("PDF 第 %d 页提取文本失败: %s", i, file_path)
            text = ""

        if not text or not text.strip():
            logger.warning("PDF 第 %d 页无文字内容，跳过: %s", i, file_path)
            continue

        documents.append(
            Document(
                page_content=text,
                metadata=_make_metadata(
                    file_path, doc_id, page=i, encoding="pdf"
                ),
            )
        )
    return documents


def load_docx_file(file_path: Path) -> list[Document]:
    """加载 DOCX 文件，将段落按合理大小分组为"页"。

    使用 python-docx 读取非空段落，按约 _DOCX_PAGE_CHARS 字符合并。

    Parameters
    ----------
    file_path : Path
        .docx 文件路径

    Returns
    -------
    list[Document]
    """
    import docx as _docx

    doc_id = _make_document_id(file_path)
    try:
        doc_obj = _docx.Document(str(file_path))
    except Exception:
        logger.warning("无法打开 DOCX 文件: %s", file_path, exc_info=True)
        return []

    paragraphs = [p.text for p in doc_obj.paragraphs if p.text and p.text.strip()]
    if not paragraphs:
        return []

    documents: list[Document] = []
    current_page = 1
    current_chars = 0
    buffer: list[str] = []

    for para in paragraphs:
        buffer.append(para)
        current_chars += len(para)
        if current_chars >= _DOCX_PAGE_CHARS:
            documents.append(
                Document(
                    page_content="\n".join(buffer),
                    metadata=_make_metadata(
                        file_path, doc_id, page=current_page, encoding="docx"
                    ),
                )
            )
            buffer.clear()
            current_chars = 0
            current_page += 1

    # 剩余段落
    if buffer:
        documents.append(
            Document(
                page_content="\n".join(buffer),
                metadata=_make_metadata(
                    file_path, doc_id, page=current_page, encoding="docx"
                ),
            )
        )

    return documents


def load_xlsx_file(file_path: Path) -> list[Document]:
    """加载 XLSX 文件，将每个工作表作为一个 Document。

    使用 openpyxl 读取所有工作表，每行转为可读文本。

    Parameters
    ----------
    file_path : Path
        .xlsx 文件路径

    Returns
    -------
    list[Document]
    """
    import openpyxl

    doc_id = _make_document_id(file_path)
    documents: list[Document] = []

    try:
        workbook = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    except Exception:
        logger.warning("无法打开 XLSX 文件: %s", file_path, exc_info=True)
        return documents

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            # 将每行的单元格值转为字符串，跳过空行
            row_values = [str(cell) if cell is not None else "" for cell in row]
            row_text = "\t".join(row_values).strip()
            if row_text:
                rows.append(row_text)

        if not rows:
            continue

        content = "\n".join(rows)
        documents.append(
            Document(
                page_content=content,
                metadata=_make_metadata(
                    file_path,
                    f"{doc_id}_{sheet_name}",
                    page=sheet_name,  # type: ignore[arg-type]
                    encoding="xlsx",
                ),
            )
        )

    workbook.close()
    return documents


def load_single_file(file_path: Path) -> list[Document]:
    """根据文件扩展名自动选择加载器。

    Parameters
    ----------
    file_path : Path
        文件路径

    Returns
    -------
    list[Document]

    Raises
    ------
    FileNotFoundError
        文件不存在
    ValueError
        不支持的文件类型
    """
    if not file_path.is_file():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = file_path.suffix.lower()
    if suffix not in _SUPPORTED_SUFFIXES:
        raise ValueError(
            f"不支持的文件类型 '{suffix}'，当前支持: {_SUPPORTED_SUFFIXES}"
        )

    # 判断知识来源类型
    knowledge_source = _detect_knowledge_source(file_path)

    if suffix == ".pdf":
        docs = load_pdf_file(file_path)
    elif suffix == ".docx":
        docs = load_docx_file(file_path)
    elif suffix == ".xlsx":
        docs = load_xlsx_file(file_path)
    else:
        # .txt / .md
        docs = load_txt_file(file_path)

    # 为所有 document 添加知识来源 metadata
    for doc in docs:
        doc.metadata["knowledge_source"] = knowledge_source

    return docs


def load_documents(directory: Path | None = None) -> list[Document]:
    """加载所有支持的知识库文件。

    默认从 data/builtin 和 data/uploads 加载。
    单个文件失败不影响其他文件，输出简短中文警告。

    Parameters
    ----------
    directory : Path, optional
        指定扫描目录；为 None 时扫描 builtin + uploads

    Returns
    -------
    list[Document]

    Raises
    ------
    RuntimeError
        未找到任何知识文件
    """
    # directory 为 None 时使用 find_knowledge_files 的默认行为（builtin + uploads）
    files = find_knowledge_files(directory)

    if not files:
        raise RuntimeError(f"未找到任何支持的知识文件（目录: {directory}）")

    all_docs: list[Document] = []
    for file_path in files:
        try:
            docs = load_single_file(file_path)
            all_docs.extend(docs)
        except Exception:
            logger.warning("加载文件失败（已跳过）: %s", file_path.name, exc_info=True)

    return all_docs


# ---------------------------------------------------------------------------
# 内部辅助
# ---------------------------------------------------------------------------


def _is_supported(path: Path) -> bool:
    """判断文件是否为支持的知识库文件（排除隐藏/临时文件）。"""
    name = path.name
    # 排除隐藏文件和临时文件
    if _EXCLUDE_PATTERNS.search(name):
        return False
    return path.suffix.lower() in _SUPPORTED_SUFFIXES


def _detect_knowledge_source(file_path: Path) -> str:
    """根据文件路径判断知识来源类型。

    Returns
    -------
    str
        "upload"
    """
    return "upload"


def _read_text_with_encoding(file_path: Path) -> tuple[str, str]:
    """依次尝试多种编码读取文件，返回 (内容, 实际编码)。"""
    raw = file_path.read_bytes()
    for enc in _TXT_ENCODINGS:
        try:
            return raw.decode(enc), enc
        except (UnicodeDecodeError, LookupError):
            continue
    # 全部失败，使用 errors="replace" 兜底
    return raw.decode("utf-8", errors="replace"), "utf-8(replace)"


def _make_document_id(file_path: Path) -> str:
    """根据文件路径生成稳定的 document_id。

    使用相对于 DATA_DIR 的路径；如果不在 DATA_DIR 下则使用文件名。

    示例：builtin/成都市煜见科技有限公司完整介绍.txt → builtin_成都市煜见科技有限公司完整介绍_txt
    """
    # 尝试相对于 DATA_DIR
    try:
        rel = file_path.resolve().relative_to(DATA_DIR.resolve())
    except ValueError:
        rel = file_path

    raw = str(rel.as_posix())
    # 替换路径分隔符和特殊字符，保留中文、英文、数字、下划线
    safe = re.sub(r"[^\w一-鿿]", "_", raw)
    # 合并多余下划线
    safe = re.sub(r"_+", "_", safe).strip("_")
    return safe or "unknown"


def _make_metadata(
    file_path: Path,
    document_id: str,
    page: int = 1,
    encoding: str = "utf-8",
) -> dict:
    """构建统一的 metadata 字典。"""
    try:
        relative_source = str(file_path.resolve().relative_to(DATA_DIR.resolve()).as_posix())
    except ValueError:
        relative_source = file_path.name
    return {
        "source": str(file_path),
        "relative_source": relative_source,
        "file_name": file_path.name,
        "file_type": file_path.suffix.lower(),
        "page": page,
        "document_id": document_id,
        "encoding": encoding,
    }
